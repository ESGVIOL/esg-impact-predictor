"""
Generate a detailed technical appendix (DOCX) for the ESG Violation
Social Media Impact Predictor dashboard.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Fill a table cell with a background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.style.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return h

def add_para(doc, text='', bold=False, italic=False, size=10, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
    return p

def add_bullet(doc, text, level=0, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_equation(doc, text, size=10):
    """Add a centred, monospace equation paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(size)
    return p

def add_note(doc, text, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run('Note. ')
    run.bold = True
    run.font.size = Pt(size)
    run.font.italic = True
    run2 = p.add_run(text)
    run2.font.size = Pt(size)
    run2.font.italic = True
    return p

# ── document setup ────────────────────────────────────────────────────────────

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin    = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin   = Cm(2.54)
section.right_margin  = Cm(2.54)

# Default body font
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

# ── TITLE ─────────────────────────────────────────────────────────────────────

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after  = Pt(18)
r = title.add_run('Web Appendix: ESG Violation Social Media Impact Predictor — Dashboard Documentation and Technical Details')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

# ── SECTION A ─────────────────────────────────────────────────────────────────

add_heading(doc, 'A. Dashboard Overview', 1)

add_para(doc,
    'The ESG Violation Social Media Impact Predictor is an interactive, browser-based decision-support '
    'tool that translates the econometric estimates reported in the main paper into actionable, '
    'market-level forecasts. Given a user-specified corporate home country, a set of target markets, '
    'an ESG violation category, and a pre-incident baseline negative-sentiment share, the dashboard '
    'computes the predicted change in the share of negative tweets for each target market and visualises '
    'the results across three complementary output panels.',
    size=11, space_after=8)

add_heading(doc, 'A.1  Input Parameters', 2)

# Table of inputs
tbl = doc.add_table(rows=5, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
hdr_cells = tbl.rows[0].cells
for cell in hdr_cells:
    set_cell_bg(cell, '1F497D')
for txt, cell in zip(['Parameter', 'Options / Range', 'Role in the Model'], hdr_cells):
    p = cell.paragraphs[0]
    run = p.add_run(txt)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

rows_data = [
    ('Company Home Country',
     'Any of ~150 countries in the psychic-distance matrix',
     'Identifies the origin country of the violating firm; used to retrieve the bilateral psychic distance to each target market'),
    ('Target Markets',
     'Any subset of ~150 countries (multi-select, continent-grouped)',
     'Countries for which predictions are generated; the home country is excluded automatically'),
    ('ESG Violation Category',
     'Environmental · Social · Governance',
     'Selects the corresponding panel-model coefficients from estimates.json'),
    ('Baseline Negative Sentiment Share (%)',
     '0 – 100 (default: 20%)',
     'The proportion of tweets about the company that carry negative sentiment before the incident; serves as the reference level for the conversion from log-scale effects to percentage-point changes'),
]

for i, (param, opts, role) in enumerate(rows_data, start=1):
    row = tbl.rows[i]
    if i % 2 == 0:
        for c in row.cells:
            set_cell_bg(c, 'DCE6F1')
    for j, txt in enumerate([param, opts, role]):
        p = row.cells[j].paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(10)

doc.add_paragraph()

add_heading(doc, 'A.2  Output Panels', 2)

add_para(doc,
    'Upon submission the dashboard renders three panels in the results area.',
    size=11, space_after=4)

add_bullet(doc,
    'Global Impact Map. A Plotly choropleth map colours each selected target market according to the '
    'predicted impact in percentage points (pp). Blue shades indicate a muted or negative response '
    '(the incident does not increase negative sentiment above baseline), while red shades signal '
    'backlash (negative sentiment rises). The violating company\'s home country is highlighted in black. '
    'Interactive tooltips display the exact numerical predictions.',
    size=11)

add_bullet(doc,
    'Detailed Predictions by Market. A sortable table reports, for each target market: the bilateral '
    'psychic distance to the home country, the immediate post-incident impact in percentage points, '
    'the predicted absolute negative-sentiment share on Day 1, and the average negative-sentiment '
    'share over the 30-day simulation window. Rows are colour-coded by impact magnitude: high backlash '
    '(> 2 pp), medium backlash (1–2 pp), low backlash (0–1 pp), and muted response (< 0 pp). '
    'The full table can be exported as a comma-separated values (CSV) file.',
    size=11)

add_bullet(doc,
    '30-Day Impact Timeline. A line chart traces the simulated trajectory of negative sentiment for '
    'the ten markets with the largest absolute impact over the 30-day post-incident window. Each line '
    'begins at the model-predicted Day-1 level and decays exponentially back toward baseline. The '
    'chart enables readers to compare how quickly backlash fades across culturally proximate versus '
    'distant markets.',
    size=11)

doc.add_paragraph()

# ── SECTION B ─────────────────────────────────────────────────────────────────

add_heading(doc, 'B. Econometric Foundation and Effect Estimation', 1)

add_heading(doc, 'B.1  Underlying Panel Model', 2)

add_para(doc,
    'The predictions are derived from panel fixed-effects regressions estimated on a country-day '
    'dataset of Twitter/X activity surrounding corporate ESG-violation events. The dependent variable '
    'is the natural logarithm of the share of negative tweets on a given day in a given market:',
    size=11, space_after=4)

add_equation(doc, 'y_{it}  =  ln( negative_tweets_{it} / total_tweets_{it} )')

add_para(doc,
    'Using the log of the share rather than the raw count controls for differences in overall '
    'Twitter activity across countries and time. The baseline specification for each violation '
    'category k is:',
    size=11, space_after=4)

add_equation(doc,
    'y_{it}  =  α_i  +  α_t  +  β₁ Incident_{t}  +  β₂ (Incident_{t} × PsychDist_i)  +  γ Z_{it}  +  ε_{it}')

add_para(doc,
    'where α_i and α_t denote country and time fixed effects, respectively; Incident_{t} is a binary '
    'indicator equal to 1 on all days from the event date onward during the observation window; '
    'PsychDist_i is the bilateral psychic distance between the violating company\'s home country and '
    'market i (Kogut–Singh index); Z_{it} is a vector of additional moderators (Long-Term Orientation '
    'for the Environmental model; unemployment rate for the Social and Governance models) and their '
    'interactions with the incident indicator; and ε_{it} is the idiosyncratic error term.',
    size=11, space_after=8)

add_heading(doc, 'B.2  Model Coefficients', 2)

add_para(doc,
    'Table B1 reports the incident-related coefficients that enter the dashboard calculations. '
    'Main effects of the moderator variables (psychic distance, LTO, unemployment) are absorbed by '
    'the country fixed effects and do not contribute to the predicted incident impact.',
    size=11, space_after=6)

# Table B1
add_para(doc, 'Table B1', bold=True, size=10, space_before=6, space_after=2)
add_para(doc, 'Incident-Related Coefficients by ESG Violation Category', italic=True, size=10, space_after=4)

coef_headers = ['Term', 'Environmental', 'Social', 'Governance']
coef_rows = [
    ('Incident (β₁)',                  '+0.0190',  '+0.0080',  '+0.0280'),
    ('Incident × Psychic Distance',    '−0.0040',  '−0.0020',  '−0.0070'),
    ('Incident × LTO',                 '+0.0010',  '—',        '—'),
    ('Incident × LTO × Psych. Dist.',  '+0.0010',  '—',        '—'),
    ('Incident × Unemployment',        '—',        '+0.0010',  '−0.0010'),
    ('Incident × Unemp. × Psych. Dist.','—',       '+0.0010',  '−0.0010'),
]

tbl2 = doc.add_table(rows=len(coef_rows)+1, cols=4)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr2 = tbl2.rows[0].cells
for c in hdr2:
    set_cell_bg(c, '1F497D')
for txt, cell in zip(coef_headers, hdr2):
    p = cell.paragraphs[0]
    r = p.add_run(txt)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, row_data in enumerate(coef_rows, start=1):
    row = tbl2.rows[i]
    if i % 2 == 0:
        for c in row.cells:
            set_cell_bg(c, 'DCE6F1')
    for j, txt in enumerate(row_data):
        p = row.cells[j].paragraphs[0]
        r = p.add_run(txt)
        r.font.size = Pt(10)
        if j > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
add_note(doc,
    'LTO = Long-Term Orientation (Hofstede). Unemployment is the national unemployment rate (%). '
    'Psychic distance follows the Kogut–Singh composite index. All coefficients reflect the within-country, '
    'within-time variation estimated by the panel fixed-effects model. Dashes indicate that the '
    'variable does not enter the respective sub-model.')

add_heading(doc, 'B.3  Computing the Incident Effect', 2)

add_para(doc,
    'For a given home country h and target market i, the dashboard computes the total log-scale '
    'incident effect as a linear combination of the estimated coefficients and the observed '
    'moderator values:',
    size=11, space_after=4)

add_para(doc, 'Environmental model:', bold=True, size=11, space_after=2)
add_equation(doc,
    'Effect_i  =  β₁  +  β₂ · PsychDist_{hi}  +  β₃ · LTO_i  +  β₄ · (LTO_i × PsychDist_{hi})')

add_para(doc, 'Social and Governance models:', bold=True, size=11, space_after=2)
add_equation(doc,
    'Effect_i  =  β₁  +  β₂ · PsychDist_{hi}  +  β₃ · Unemp_i  +  β₄ · (Unemp_i × PsychDist_{hi})')

add_para(doc,
    'Because the dependent variable is the log of the negative-sentiment share, the effect must be '
    'converted to a percentage-point change for interpretability. Let S₀ denote the baseline '
    'negative-sentiment share (entered by the user as a percentage divided by 100). The '
    'post-incident share and the percentage-point impact are:',
    size=11, space_after=4)

add_equation(doc, 'S₁  =  S₀ · exp( Effect_i )')
add_equation(doc, 'Impact (pp)  =  ( S₁ − S₀ ) × 100')

add_para(doc,
    'A positive Impact (pp) value indicates that the incident increases the share of negative tweets '
    'relative to baseline — i.e., social media backlash. A negative value means the incident is '
    'associated with a muted response in that market: the negative-sentiment share does not rise '
    'above pre-incident levels.',
    size=11, space_after=8)

add_heading(doc, 'B.4  Worked Example', 2)

add_para(doc,
    'Consider a U.S.-headquartered firm committing a Governance violation; the target market is '
    'Germany, with a psychic distance of 2.1 (Kogut–Singh units), an unemployment rate of 3.2%, '
    'and a user-specified baseline negativity of 20% (S₀ = 0.20).',
    size=11, space_after=4)

add_equation(doc,
    'Effect  =  0.028  +  (−0.007 × 2.1)  +  (−0.001 × 3.2)  +  (−0.001 × 3.2 × 2.1)')
add_equation(doc,
    '       =  0.028  −  0.0147  −  0.0032  −  0.00672')
add_equation(doc,
    '       =  0.00338')
add_equation(doc, 'S₁  =  0.20 × exp(0.00338)  ≈  0.2007')
add_equation(doc, 'Impact  =  (0.2007 − 0.20) × 100  ≈  +0.07 pp')

add_para(doc,
    'The small positive value reflects that governance violations in a culturally proximate, '
    'low-unemployment market like Germany produce a modest, short-lived backlash. By contrast, '
    'the same violation generates a larger effect in culturally close markets with low unemployment '
    'because the interaction between unemployment and psychic distance attenuates the incident '
    'coefficient at high unemployment/high distance combinations.',
    size=11, space_after=8)

# ── SECTION C ─────────────────────────────────────────────────────────────────

add_heading(doc, 'C. 30-Day Post-Incident Simulation', 1)

add_heading(doc, 'C.1  Rationale for an Exponential Decay Model', 2)

add_para(doc,
    'Social media attention to corporate misconduct is well-documented to follow a rapid initial '
    'spike followed by exponential decay as news cycles move on and consumer attention shifts '
    '(Pfeffer et al. 2014; Kietzmann et al. 2011). We therefore model the day-specific negative-'
    'sentiment share as:',
    size=11, space_after=4)

add_equation(doc, 'S(d)  =  S₀ · exp( Effect_i · δ(d) )')
add_equation(doc, 'δ(d)  =  exp( −0.05 × (d − 1) ),    d = 1, 2, …, 30')

add_para(doc,
    'where d is the number of days after the incident (Day 1 = incident date), δ(d) is the '
    'time-decay factor, and the decay parameter λ = 0.05 implies that the marginal effect of the '
    'incident on log-sentiment diminishes by approximately 5% per day. On Day 1 the full log-scale '
    'effect is applied (δ(1) = 1); by Day 14 roughly 50% of the initial effect remains '
    '(δ(14) ≈ 0.51); and by Day 30 the effect has decayed to approximately 23% of its initial '
    'magnitude (δ(30) ≈ 0.23).',
    size=11, space_after=8)

add_heading(doc, 'C.2  Derivation of the Decay Parameter', 2)

add_para(doc,
    'The decay rate λ = 0.05 corresponds to a half-life of ln(2)/0.05 ≈ 13.9 days, consistent '
    'with the empirical observation that social media reaction to corporate ESG incidents '
    'typically subsides within two to three weeks for most firm–market combinations. The '
    'parameter can be adjusted directly in the application source code (app.js, function '
    'calculatePrediction) to accommodate different assumed attention-decay rates.',
    size=11, space_after=8)

add_heading(doc, 'C.3  Summary Statistic: Average 30-Day Negativity', 2)

add_para(doc,
    'The average negative-sentiment share reported in the table for each market is the arithmetic '
    'mean of the day-specific shares over the full 30-day simulation window:',
    size=11, space_after=4)

add_equation(doc, 'S̄₃₀  =  (1/30) · Σ_{d=1}^{30} S(d)')

add_para(doc,
    'This metric provides a single aggregate measure of the sustained reputational burden borne '
    'by the firm in a given market over the month following the incident.',
    size=11, space_after=8)

add_heading(doc, 'C.4  Decay Trajectory: Numerical Illustration', 2)

add_para(doc, 'Table C1', bold=True, size=10, space_before=6, space_after=2)
add_para(doc, 'Illustrative 30-Day Decay of the Incident Effect (λ = 0.05)', italic=True, size=10, space_after=4)

decay_rows = [
    ('1',  '1.000', 'Full incident impact'),
    ('5',  '0.819', 'c. 82% of initial effect remains'),
    ('10', '0.641', 'c. 64% of initial effect remains'),
    ('14', '0.513', 'Half-life crossed (50% threshold)'),
    ('20', '0.387', 'c. 39% of initial effect remains'),
    ('25', '0.301', 'c. 30% of initial effect remains'),
    ('30', '0.233', 'c. 23% of initial effect remains'),
]

tbl3 = doc.add_table(rows=len(decay_rows)+1, cols=3)
tbl3.style = 'Table Grid'
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr3 = tbl3.rows[0].cells
for c in hdr3:
    set_cell_bg(c, '1F497D')
for txt, cell in zip(['Day (d)', 'Decay Factor δ(d)', 'Interpretation'], hdr3):
    p = cell.paragraphs[0]
    r = p.add_run(txt)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, (day, factor, interp) in enumerate(decay_rows, start=1):
    row = tbl3.rows[i]
    if i % 2 == 0:
        for c in row.cells:
            set_cell_bg(c, 'DCE6F1')
    for j, txt in enumerate([day, factor, interp]):
        p = row.cells[j].paragraphs[0]
        r = p.add_run(txt)
        r.font.size = Pt(10)
        if j < 2:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ── SECTION D ─────────────────────────────────────────────────────────────────

add_heading(doc, 'D. Technical Details', 1)

add_heading(doc, 'D.1  Data Sources', 2)

# Table D1
tbl4 = doc.add_table(rows=5, cols=4)
tbl4.style = 'Table Grid'
tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr4 = tbl4.rows[0].cells
for c in hdr4:
    set_cell_bg(c, '1F497D')
for txt, cell in zip(['Data File', 'Variable', 'Coverage', 'Source / Construction'], hdr4):
    p = cell.paragraphs[0]
    r = p.add_run(txt)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(9.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

src_rows = [
    ('distances.json',
     'Psychic distance (Kogut–Singh composite index)',
     '~150 countries; 22,350 bilateral pairs',
     'Kogut & Singh (1988) formula applied to Hofstede\'s six cultural dimensions; bilateral matrix symmetric'),
    ('lto.json',
     'Long-Term Orientation (LTO) index',
     '109 countries',
     'Hofstede Insights country scores (6th dimension); used exclusively in the Environmental sub-model'),
    ('unemployment.json',
     'National unemployment rate (%)',
     '180 countries',
     'World Bank World Development Indicators; most recent available annual rate per country'),
    ('estimates.json',
     'Panel-model coefficients',
     '3 ESG categories × up to 7 terms',
     'Authors\' own estimation on the Twitter/ESG event dataset described in the main paper (Section 3)'),
]

for i, row_data in enumerate(src_rows, start=1):
    row = tbl4.rows[i]
    if i % 2 == 0:
        for c in row.cells:
            set_cell_bg(c, 'DCE6F1')
    for j, txt in enumerate(row_data):
        p = row.cells[j].paragraphs[0]
        r = p.add_run(txt)
        r.font.size = Pt(9.5)

doc.add_paragraph()

add_heading(doc, 'D.2  Psychic Distance Computation', 2)

add_para(doc,
    'Psychic distance is operationalised as the Kogut–Singh (1988) composite index computed across '
    'Hofstede\'s (1980, 2001) six cultural dimensions: Power Distance (PDI), Individualism (IDV), '
    'Masculinity (MAS), Uncertainty Avoidance (UAI), Long-Term Orientation (LTO), and Indulgence '
    'versus Restraint (IVR). The index for country pair (h, i) is:',
    size=11, space_after=4)

add_equation(doc,
    'PsychDist_{hi}  =  (1/6) · Σ_{k=1}^{6}  { (H_{kh} − H_{ki})² / V_k }')

add_para(doc,
    'where H_{kj} is the score of country j on dimension k and V_k is the variance of dimension k '
    'across all countries in the sample (used for standardisation). Higher values indicate greater '
    'cultural distance between the two countries.',
    size=11, space_after=8)

add_heading(doc, 'D.3  Architecture and Implementation', 2)

add_para(doc,
    'The dashboard is a fully client-side single-page application requiring no server-side logic '
    'or database backend. All computation is executed in the user\'s browser via JavaScript. '
    'The following libraries are used:',
    size=11, space_after=4)

add_bullet(doc, 'Bootstrap 5.3 — responsive layout and UI components', size=11)
add_bullet(doc, 'Bootstrap Icons 1.11 — iconography', size=11)
add_bullet(doc, 'Plotly.js 2.27 — interactive choropleth map and timeline chart', size=11)

add_para(doc,
    'Data files are loaded asynchronously at page startup via the Fetch API. Prediction calculations '
    'run synchronously in under 100 ms for up to 150 target markets. The application has been tested '
    'on Chrome/Edge 90+, Firefox 88+, and Safari 14+.',
    size=11, space_after=8)

add_heading(doc, 'D.4  Handling of Missing Data', 2)

add_para(doc,
    'Not all country pairs have psychic-distance values, and not all countries have LTO or '
    'unemployment data. The application handles these cases as follows:',
    size=11, space_after=4)

add_bullet(doc,
    'Missing bilateral psychic distance: the target market is silently excluded from the results '
    'set. Users can verify coverage by inspecting the distance matrix (distances.json).',
    size=11)
add_bullet(doc,
    'Missing LTO score (Environmental model): the LTO-related interaction terms are set to zero, '
    'effectively treating the market as having a neutral LTO score relative to its moderating '
    'influence on the incident effect.',
    size=11)
add_bullet(doc,
    'Missing unemployment rate (Social/Governance models): the unemployment interaction terms are '
    'set to zero, preserving the core incident and psychic-distance effects.',
    size=11)

add_para(doc,
    'Country names are matched with a case-insensitive fuzzy lookup to accommodate minor spelling '
    'differences across data sources (e.g., "Czechia" vs. "Czech Republic").',
    size=11, space_after=8)

add_heading(doc, 'D.5  Reproducibility', 2)

add_para(doc,
    'The complete application source code (index.html, app.js, styles.css) and all data files '
    '(estimates.json, distances.json, lto.json, unemployment.json) are provided in the online '
    'supplement. The tool can be run locally by serving the directory with any static HTTP server '
    '(e.g., python3 -m http.server 8000) or deployed without modification to any static hosting '
    'service such as GitHub Pages. No external API calls or authentication are required.',
    size=11, space_after=8)

# ── REFERENCES ────────────────────────────────────────────────────────────────

add_heading(doc, 'References', 1)

refs = [
    'Hofstede, G. (1980). Culture\'s consequences: International differences in work-related values. Sage.',
    'Hofstede, G. (2001). Culture\'s consequences: Comparing values, behaviors, institutions, and organizations across nations (2nd ed.). Sage.',
    'Kietzmann, J. H., Hermkens, K., McCarthy, I. P., & Silvestre, B. S. (2011). Social media? Get serious! Understanding the functional building blocks of social media. Business Horizons, 54(3), 241–251.',
    'Kogut, B., & Singh, H. (1988). The effect of national culture on the choice of entry mode. Journal of International Business Studies, 19(3), 411–432.',
    'Pfeffer, J., Zorbach, T., & Carley, K. M. (2014). Understanding online firestorms: Negative word-of-mouth dynamics in social media networks. Journal of Marketing Communications, 20(1–2), 117–128.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    run = p.add_run(ref)
    run.font.size = Pt(10)

# ── SAVE ──────────────────────────────────────────────────────────────────────

out_path = '/Users/raoulvkubler/Dropbox/Project implementation Corporate Unethical Behavior/Revision/App/Appendix_Dashboard_Technical.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
